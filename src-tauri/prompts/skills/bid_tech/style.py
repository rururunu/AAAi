# -*- coding: utf-8 -*-
"""技术标版式预设：页边距、页眉、标题/正文字体、表格线框样式。

约定（可被调用方覆盖）：
- 一级标题：黑体、小二（18pt）、加粗
- 二级标题：黑体、小三（16pt）、加粗
- 三级标题：黑体、四号（14pt）、加粗
- 正文：仿宋、小四（12pt）；表内文字：宋体、五号（10.5pt）
- 页边距：上下左右约 2.5cm
"""

from __future__ import annotations

from typing import Iterable, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# --- 字体名（Windows 常见，避免依赖厂商品牌字体包） ---
FONT_TITLE = "黑体"
FONT_BODY = "仿宋"
FONT_TABLE = "宋体"

SIZE_H1 = 18
SIZE_H2 = 16
SIZE_H3 = 14
SIZE_BODY = 12
SIZE_TABLE = 10.5
SIZE_HEADER = 9


def set_run_font(
    run: Run,
    *,
    name: str = FONT_BODY,
    size_pt: float = SIZE_BODY,
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> Run:
    """为单个 run 同时设置西文名与东亚字体名，避免中文回落成系统默认字体。"""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    return run


def configure_document(
    doc: Optional[Document] = None,
    *,
    margin_cm: float = 2.5,
    header_text: str = "",
) -> Document:
    """创建或配置 Document：页边距 + 可选页眉。"""
    document = doc or Document()
    section = document.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)

    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_paragraph_runs(paragraph)
        run = paragraph.add_run(header_text)
        set_run_font(run, name=FONT_TABLE, size_pt=SIZE_HEADER)

    return document


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    """删除段落内全部 run（python-docx 无稳定的 clear API）。"""
    element = paragraph._element
    for child in list(element):
        if child.tag == qn("w:r"):
            element.remove(child)


def add_heading_cn(
    doc: Document,
    text: str,
    *,
    level: int = 1,
) -> Paragraph:
    """中文技术标标题。level: 1/2/3。"""
    if level <= 1:
        size, space_before, space_after = SIZE_H1, Pt(12), Pt(8)
    elif level == 2:
        size, space_before, space_after = SIZE_H2, Pt(10), Pt(6)
    else:
        size, space_before, space_after = SIZE_H3, Pt(8), Pt(4)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run, name=FONT_TITLE, size_pt=size, bold=True)
    return paragraph


def add_body(
    doc: Document,
    text: str,
    *,
    first_line_indent: bool = True,
) -> Paragraph:
    """正文段落（仿宋）。表格场景下的说明句也可走这里，但不得用正文冒充整表。"""
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(4)
    if first_line_indent:
        fmt.first_line_indent = Cm(0.74)  # 约两字符
    run = paragraph.add_run(text)
    set_run_font(run, name=FONT_BODY, size_pt=SIZE_BODY)
    return paragraph


def add_score_lead(
    doc: Document,
    chapter_title: str,
    score_points: Sequence[str],
) -> Paragraph:
    """每章开头一句对照评分点（强制约束，便于评委扫读）。"""
    joined = "；".join(score_points) if score_points else "见本章表格"
    return add_body(
        doc,
        f"【评分对照】{chapter_title}：对应评分点——{joined}。",
        first_line_indent=False,
    )


def _set_cell_border(cell: _Cell, **edges: str) -> None:
    """设置单元格边框。edges 键为 top/left/bottom/right，值为 'nil' 或如 '4'（1/8 pt）。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge_name, value in edges.items():
        element = OxmlElement(f"w:{edge_name}")
        if value == "nil":
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), value)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        # 替换同名旧边框
        for old in list(tc_borders):
            if old.tag == element.tag:
                tc_borders.remove(old)
        tc_borders.append(element)


def apply_grid_borders(table: Table, *, sz: str = "4") -> None:
    """全表格线框：单元格四边细实线（投标/公文常见样式）。"""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=sz, bottom=sz, left=sz, right=sz)


def apply_three_line_table(table: Table) -> None:
    """三线表：顶线粗、表头下细线、底线粗，无竖线。"""
    rows = table.rows
    if not rows:
        return
    last = len(rows) - 1
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            top = "12" if r_idx == 0 else ("6" if r_idx == 1 else "nil")
            bottom = "12" if r_idx == last else ("6" if r_idx == 0 else "nil")
            # 表头行（第 0 行）底边用细线；数据行无横线，靠底线收束
            if r_idx == 0:
                bottom = "6"
            elif r_idx == last:
                top = "nil"
                bottom = "12"
            else:
                top = "nil"
                bottom = "nil"
            _set_cell_border(cell, top=top, bottom=bottom, left="nil", right="nil")


def fill_cell(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    font_name: str = FONT_TABLE,
    size_pt: float = SIZE_TABLE,
    align_center: bool = False,
) -> None:
    """清空单元格并写入统一字体文本。"""
    # 保留一个段落，清掉多余段落与 runs
    paragraphs = cell.paragraphs
    if not paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = paragraphs[0]
        for extra in paragraphs[1:]:
            p = extra._element
            p.getparent().remove(p)
        _clear_paragraph_runs(paragraph)

    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text if text is not None else "")
    set_run_font(run, name=font_name, size_pt=size_pt, bold=bold)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    *,
    grid: bool = True,
    three_line: bool = False,
) -> Table:
    """创建带表头的真表格（python-docx add_table），并统一单元格字体。

    默认 ``grid=True``：全表四边线框。若需三线表，设 ``three_line=True, grid=False``。
    """
    row_list = [list(map(lambda x: "" if x is None else str(x), row)) for row in rows]
    table = doc.add_table(rows=1 + len(row_list), cols=len(headers))
    table.autofit = True

    for col, header in enumerate(headers):
        fill_cell(table.rows[0].cells[col], header, bold=True, align_center=True)

    for r_idx, row in enumerate(row_list, start=1):
        for c_idx, value in enumerate(row):
            if c_idx >= len(headers):
                break
            fill_cell(table.rows[r_idx].cells[c_idx], value)

    if three_line:
        apply_three_line_table(table)
    elif grid:
        apply_grid_borders(table)
    return table
