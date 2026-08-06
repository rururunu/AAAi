# -*- coding: utf-8 -*-
"""可复用技术标表模板。

所有模板均调用 ``style.add_table``（python-docx ``Document.add_table``），
禁止用多段文字模拟表格行列。

内置模板：
- 日程四列表：时间 / 活动内容 / 地点 / 负责人员
- 流程五列表：序号 / 工作事项 / 具体做法 / 责任岗位 / 形成材料
- 配比四列表：岗位 / 配置标准 / 资质要求 / 人数
- 档案三列表：类别 / 归档内容 / 形成或更新时点
- 架构三列表：层级 / 组成 / 职责
- 应急三列表：场景 / 响应动作 / 第一责任人
- 备用方案三列表：触发条件 / 原计划 / 备用方案
- 场馆核实四列表：场馆 / 开放与预约 / 讲解 / 雨天备选
- 保险核实五列表：险种 / 保额 / 出单时效 / 理赔热线 / 备注
- 交通核实四列表：航段或路段 / 建议方式 / 集散节点 / 备注
"""

from __future__ import annotations

from typing import Iterable, Sequence

from docx.document import Document
from docx.table import Table

from . import style

# --- 表头常量（集中管理，避免各脚本拼写漂移） ---
HEADERS_SCHEDULE = ("时间", "活动内容", "地点", "负责人员")
HEADERS_PROCESS = ("序号", "工作事项", "具体做法", "责任岗位", "形成材料")
HEADERS_STAFFING = ("岗位", "配置标准", "资质要求", "人数")
HEADERS_ARCHIVE = ("类别", "归档内容", "形成/更新时点")
HEADERS_ORG = ("层级", "组成", "职责")
HEADERS_EMERGENCY = ("场景", "响应动作", "第一责任人")
HEADERS_BACKUP = ("触发条件", "原计划", "备用方案")
HEADERS_VENUE = ("场馆", "开放时间与预约规则", "讲解安排", "雨天备选")
HEADERS_INSURANCE = ("险种", "保额", "出单时效", "理赔热线", "备注")
HEADERS_TRANSPORT = ("航段/路段", "建议方式", "集散节点", "备注")
HEADERS_ACCEPTANCE = ("环节", "我方配合事项", "提交材料")
HEADERS_OVERVIEW = ("天数", "地点", "主要活动内容", "重点场馆（如有）")
HEADERS_SCHEDULE_DAY = ("天数", "时间", "日程安排", "线路规划", "酒店")
HEADERS_TIMELINE = ("时间", "具体环节", "工作内容", "备注")


def add_schedule_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """半小时日程四列表。行格式: (时间, 活动内容, 地点, 负责人员)。"""
    return style.add_table(doc, HEADERS_SCHEDULE, rows)


def add_process_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """阶段流程五列表。行格式: (序号, 事项, 做法, 岗位, 形成材料)。"""
    return style.add_table(doc, HEADERS_PROCESS, rows)


def add_staffing_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """人员配比四列表。"""
    return style.add_table(doc, HEADERS_STAFFING, rows)


def add_archive_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """档案分类三列表。"""
    return style.add_table(doc, HEADERS_ARCHIVE, rows)


def add_org_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """组织架构三列表。"""
    return style.add_table(doc, HEADERS_ORG, rows)


def add_emergency_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """安全应急：场景 × 响应动作 × 第一责任人。"""
    return style.add_table(doc, HEADERS_EMERGENCY, rows)


def add_backup_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """活动备用方案表。"""
    return style.add_table(doc, HEADERS_BACKUP, rows)


def add_venue_research_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """联网核实场馆信息：只填格，不写亮点散文。"""
    return style.add_table(doc, HEADERS_VENUE, rows)


def add_insurance_research_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """联网核实保险字段。"""
    return style.add_table(doc, HEADERS_INSURANCE, rows)


def add_transport_research_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """联网核实交通字段。"""
    return style.add_table(doc, HEADERS_TRANSPORT, rows)


def add_acceptance_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """验收 / 审计配合表。"""
    return style.add_table(doc, HEADERS_ACCEPTANCE, rows)


def add_overview_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """行程总览表（活动方案必备）。"""
    return style.add_table(doc, HEADERS_OVERVIEW, rows)


def add_day_schedule_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """按天块行程五列表：天数 / 时段 / 日程 / 线路 / 酒店。"""
    return style.add_table(doc, HEADERS_SCHEDULE_DAY, rows)


def add_timeline_table(
    doc: Document,
    rows: Iterable[Sequence[str]],
) -> Table:
    """单日流程四列表：时间 / 具体环节 / 工作内容 / 备注。"""
    return style.add_table(doc, HEADERS_TIMELINE, rows)
